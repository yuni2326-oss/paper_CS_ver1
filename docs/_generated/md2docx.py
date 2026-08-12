"""간이 Markdown→docx 변환기(python-docx). 헤딩/문단/표/리스트/인용/볼드·이탤릭·코드 처리.
LaTeX 수식($...$/$$...$$)은 렌더 없이 원문 텍스트로 보존(작업용 docx). 사용:
  python md2docx.py in.md out.docx [제목]
"""
import os
import re
import sys
from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

IMG = re.compile(r'^!\[[^\]]*\]\(([^)]+)\)\s*$')

# ── 한글 폰트 처리 ────────────────────────────────────────────────────────────
# python-docx는 w:rFonts의 **w:eastAsia**를 설정하지 않는다. 그러면 Word가 한글 글리프를
# 테마 라틴 폰트(Calibri 등)에서 찾다 실패해 한글이 깨져 보인다(영문 문서에서는 드러나지 않음).
# 따라서 (1) 문서 기본값(docDefaults), (2) 사용하는 모든 스타일, (3) 폰트를 직접 지정하는
# 런(코드 런) 세 곳 모두에 eastAsia를 명시한다.
EA_FONT = '맑은 고딕'          # Malgun Gothic — 한국어 Windows 표준 탑재
LATIN_FONT = 'Calibri'
CODE_FONT = 'Consolas'
STYLES_TO_PATCH = ('Normal', 'Title', 'Heading 1', 'Heading 2', 'Heading 3', 'Heading 4',
                   'List Bullet', 'List Number', 'Intense Quote', 'Caption',
                   'Light Grid Accent 1')


def _set_rfonts(rpr, latin=LATIN_FONT, ea=EA_FONT):
    """rPr에 w:rFonts(ascii/hAnsi/eastAsia/cs)를 명시. eastAsia가 한글 렌더의 핵심.
    ⚠️ w:*Theme 속성이 남아 있으면 Word가 명시 폰트 대신 테마 폰트를 우선할 수 있으므로
    반드시 제거한다(이게 빠지면 eastAsia를 넣어도 한글이 계속 깨진다)."""
    rf = rpr.find(qn('w:rFonts'))
    if rf is None:
        rf = rpr.makeelement(qn('w:rFonts'), {})
        rpr.insert(0, rf)
    for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
        if rf.get(qn(theme_attr)) is not None:
            del rf.attrib[qn(theme_attr)]
    rf.set(qn('w:ascii'), latin)
    rf.set(qn('w:hAnsi'), latin)
    rf.set(qn('w:eastAsia'), ea)
    rf.set(qn('w:cs'), latin)


def setup_korean_fonts(doc):
    """문서 기본값 + 스타일 전반에 한글 폰트를 심는다."""
    # (1) docDefaults — 스타일이 명시하지 않은 모든 텍스트가 이걸 상속
    styles_el = doc.styles.element
    dd = styles_el.find(qn('w:docDefaults'))
    if dd is None:
        dd = styles_el.makeelement(qn('w:docDefaults'), {})
        styles_el.insert(0, dd)
    rpd = dd.find(qn('w:rPrDefault'))
    if rpd is None:
        rpd = dd.makeelement(qn('w:rPrDefault'), {})
        dd.insert(0, rpd)
    rpr = rpd.find(qn('w:rPr'))
    if rpr is None:
        rpr = rpd.makeelement(qn('w:rPr'), {})
        rpd.insert(0, rpr)
    _set_rfonts(rpr)
    lang = rpr.find(qn('w:lang'))
    if lang is None:
        lang = rpr.makeelement(qn('w:lang'), {})
        rpr.append(lang)
    lang.set(qn('w:val'), 'en-US')
    lang.set(qn('w:eastAsia'), 'ko-KR')

    # (2) 본문 계열 스타일 — 헤딩은 테마 폰트를 쓰므로 라틴까지 명시적으로 덮어쓴다.
    for name in STYLES_TO_PATCH:
        try:
            st = doc.styles[name]
        except KeyError:
            continue
        _set_rfonts(st.element.get_or_add_rPr())

    # (3) styles.xml 안의 **나머지 모든 w:rFonts** (표 스타일의 조건부 서식 w:tblStylePr 등
    #     python-docx API가 닿지 않는 곳)에는 테마 속성만 제거하고 eastAsia만 심는다.
    #     ascii/hAnsi는 건드리지 않는다 — 기호 폰트(Symbol/Wingdings)를 쓰는 스타일을
    #     망가뜨리지 않기 위함.
    for rf in styles_el.iter(qn('w:rFonts')):
        for theme_attr in ('w:asciiTheme', 'w:hAnsiTheme', 'w:eastAsiaTheme', 'w:cstheme'):
            if rf.get(qn(theme_attr)) is not None:
                del rf.attrib[qn(theme_attr)]
        rf.set(qn('w:eastAsia'), EA_FONT)


def set_run_font(run, latin, ea=EA_FONT):
    """런 단위 폰트 지정(코드 런 등) — 한글이 섞여도 깨지지 않게 eastAsia 동반 설정."""
    run.font.name = latin
    _set_rfonts(run._element.get_or_add_rPr(), latin=latin, ea=ea)

INLINE = re.compile(r'(\*\*.+?\*\*|`[^`]+`|(?<!\*)\*(?!\s)(?!\*)[^*]+?\*(?!\*))')


def add_runs(p, text):
    pos = 0
    for m in INLINE.finditer(text):
        if m.start() > pos:
            p.add_run(text[pos:m.start()])
        tok = m.group(0)
        if tok.startswith('**'):
            p.add_run(tok[2:-2]).bold = True
        elif tok.startswith('`'):
            set_run_font(p.add_run(tok[1:-1]), CODE_FONT)
        else:
            p.add_run(tok[1:-1]).italic = True
        pos = m.end()
    if pos < len(text):
        p.add_run(text[pos:])


def is_sep(line):
    s = line.strip()
    return s.startswith('|') and set(s) <= set('|-: ') and '-' in s


def row_cells(line):
    """표 한 줄을 칸으로 나눈다. **`\\|`는 칸 구분자가 아니다.**

    이스케이프를 무시하고 `split('|')`하면 셀 안의 파이프에서 칸이 쪼개져 열이 밀리고 그
    뒤 내용이 잘려 나간다(부록 B의 `|rFFT|`가 그랬다). 이스케이프된 파이프는 문자 그대로
    복원한다."""
    s = line.strip()
    if s.startswith('|'):
        s = s[1:]
    if s.endswith('|') and not s.endswith('\\|'):
        s = s[:-1]
    out, cur, i = [], [], 0
    while i < len(s):
        if s[i] == '\\' and i + 1 < len(s) and s[i + 1] == '|':
            cur.append('|'); i += 2
        elif s[i] == '|':
            out.append(''.join(cur).strip()); cur = []; i += 1
        else:
            cur.append(s[i]); i += 1
    out.append(''.join(cur).strip())
    return out


# HTML 주석은 Word에 **보이는 텍스트**로 새어 나온다. 마크다운에서 주석은 정의상
# 출력에 나오지 않아야 하므로 변환 전에 제거한다. 논문 2가 표·그림 자리표시자로
# `<!-- AUTO:table_1 -->`를 쓰면서 제출본에 그대로 찍히는 것이 발견됐다.
HTML_COMMENT = re.compile(r'<!--.*?-->', re.DOTALL)


def _trPr(row):
    """행 속성 요소를 얻거나 만든다."""
    tr = row._tr
    pr = tr.find(qn('w:trPr'))
    if pr is None:
        pr = OxmlElement('w:trPr')
        tr.insert(0, pr)
    return pr


def _repeat_header(row):
    """이 행을 페이지마다 반복되는 머리행으로 표시한다."""
    pr = _trPr(row)
    if pr.find(qn('w:tblHeader')) is None:
        el = OxmlElement('w:tblHeader')
        el.set(qn('w:val'), 'true')
        pr.append(el)


def _no_split(row):
    """행이 페이지 경계에서 쪼개지지 않게 한다."""
    pr = _trPr(row)
    if pr.find(qn('w:cantSplit')) is None:
        pr.append(OxmlElement('w:cantSplit'))


def main(inp, outp, title=None):
    base = os.path.dirname(os.path.abspath(inp))
    raw = open(inp, encoding='utf-8').read()
    lines = [ln for ln in HTML_COMMENT.sub('', raw).splitlines()]
    doc = Document()
    doc.styles['Normal'].font.size = Pt(10.5)
    setup_korean_fonts(doc)   # 한글 깨짐 방지(eastAsia 폰트)
    i, n = 0, len(lines)
    while i < n:
        line = lines[i]
        s = line.strip()
        # 이미지: ![alt](path) → 그림 삽입(경로는 md 위치 기준 상대해석)
        mimg = IMG.match(s)
        if mimg:
            p = mimg.group(1)
            fp = p if os.path.isabs(p) else os.path.join(base, p)
            para = doc.add_paragraph(); para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            try:
                para.add_run().add_picture(fp, width=Inches(6.0))
            except Exception as e:
                para.add_run(f"[figure missing: {p} ({type(e).__name__})]")
            i += 1; continue
        # 표: 헤더행 + 구분행
        if s.startswith('|') and i + 1 < n and is_sep(lines[i + 1]):
            header = row_cells(line)
            rows = []
            i += 2
            while i < n and lines[i].strip().startswith('|'):
                rows.append(row_cells(lines[i])); i += 1
            t = doc.add_table(rows=1, cols=len(header)); t.style = 'Light Grid Accent 1'
            # **큰 표가 페이지를 넘어갈 때** 머리행이 반복되지 않고 행이 중간에서 잘려
            # 읽을 수 없게 되던 것을 고친다(Table 2·11·부록 B가 그랬다).
            #   tblHeader  = 첫 행을 각 페이지에서 반복
            #   cantSplit  = 한 행이 페이지 경계에서 쪼개지지 않게
            _repeat_header(t.rows[0])
            for r_ in t.rows:
                _no_split(r_)
            for j, h in enumerate(header):
                add_runs(t.rows[0].cells[j].paragraphs[0], h)
                for r in t.rows[0].cells[j].paragraphs[0].runs: r.bold = True
            for r in rows:
                row_obj = t.add_row()
                _no_split(row_obj)
                cells = row_obj.cells
                for j in range(min(len(r), len(header))):
                    add_runs(cells[j].paragraphs[0], r[j])
            doc.add_paragraph()
            continue
        if not s:
            i += 1; continue
        # 헤딩
        m = re.match(r'^(#{1,6})\s+(.*)', s)
        if m:
            lvl = len(m.group(1))
            h = doc.add_heading(level=min(lvl, 4))
            add_runs(h, re.sub(r'\s*#+\s*$', '', m.group(2)))
            i += 1; continue
        # 수평선
        if re.match(r'^([-*_])\1{2,}$', s):
            i += 1; continue
        # 인용
        if s.startswith('>'):
            p = doc.add_paragraph(style='Intense Quote'); add_runs(p, s.lstrip('> ').strip())
            i += 1; continue
        # 리스트(번호/불릿)
        mnum = re.match(r'^(\d+)\.\s+(.*)', s)
        mbul = re.match(r'^[-*]\s+(.*)', s)
        if mnum:
            # **번호를 문자 그대로 싣는다.** `List Number` 스타일을 쓰면 Word가 번호를
            # 자동생성하고 문서 전체에서 하나의 목록으로 이어버린다 — 논문 2에서
            # §5.1의 "다섯 근거"가 5–9로, §6.2가 10–17로, 결론이 18–21로 찍혔다.
            # 저자가 쓴 번호가 정본이므로 그것을 그대로 쓰고 내어쓰기만 준다.
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Inches(0.35)
            p.paragraph_format.first_line_indent = Inches(-0.25)
            add_runs(p, f"{mnum.group(1)}. {mnum.group(2)}")
            i += 1; continue
        if mbul:
            p = doc.add_paragraph(style='List Bullet'); add_runs(p, mbul.group(1)); i += 1; continue
        # 일반 문단
        p = doc.add_paragraph(); add_runs(p, s)
        i += 1
    if title:
        core = doc.core_properties; core.title = title
    doc.save(outp)

    # 자체검증(회귀 방지): 한글 문서가 eastAsia 폰트 없이 저장되면 Word에서 깨진다.
    import zipfile
    st = zipfile.ZipFile(outp).read('word/styles.xml').decode('utf-8')
    n_theme = len(re.findall(r'w:(?:ascii|hAnsi|eastAsia|cs)Theme=', st))
    n_ea = st.count(f'w:eastAsia="{EA_FONT}"')
    if n_ea == 0 or n_theme > 0:
        raise AssertionError(
            f"한글 폰트 설정 실패: eastAsia={n_ea}개, 남은 theme 속성={n_theme}개 "
            f"(theme 속성이 남으면 Word가 명시 폰트를 무시할 수 있음)")
    has_kr = any('가' <= c <= '힣' for c in open(inp, encoding='utf-8').read())
    print(f"OK: {inp} -> {outp} ({len(lines)} lines)"
          f"{f' | 한글 폰트 적용({EA_FONT}), eastAsia {n_ea}곳' if has_kr else ''}")


if __name__ == "__main__":
    main(sys.argv[1], sys.argv[2], sys.argv[3] if len(sys.argv) > 3 else None)
